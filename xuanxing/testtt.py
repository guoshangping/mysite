import docx
from docx import Document
import re

def iter_headings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph


def title_deal(file_mulu):
    doc = docx.Document(file_mulu)
    t1 = 0
    t2 = 0
    t3 = 0
    t4 = 0
    p1 = re.compile(r'\d.*\.\d+', re.S)  # 替换该正则匹配到的字符串
    title_list = []
    for heading in iter_headings(doc.paragraphs):
        if heading.style.name == "Heading 1":
            t1 += 1
            t2 = 0
            t3 = 0
            t4 = 0
        if heading.style.name == 'Heading 2':
            t2 = t2 + 1
            biaoti = str(t1) + "." + str(t2)
            t3 = 0
        elif heading.style.name == 'Heading 3':
            t3 = t3 + 1
            biaoti = str(t1) + "." + str(t2) + "." + str(t3)
            t4 = 0
        elif heading.style.name == 'Heading 4':
            t4 = t4 + 1
            biaoti = str(t1) + "." + str(t2) + "." + str(t3) + "." + str(t4)
        else:
            continue
        if not re.findall(p1, heading.text):
            heading.text = biaoti + heading.text
        else:
            heading.text = re.sub(p1, biaoti, heading.text)
        title_list.append(heading.text)
    title_list = [tt.strip().replace("\n", "") for tt in title_list]
    return title_list


if __name__ == '__main__':
    print("dsada")
# if __name__ == '__main__':
#     fn = "./t5.docx"
#     doc = docx.Document(fn)
#     print(type(doc.paragraphs))
#     docs = [pj for pj in doc.paragraphs if pj.text.strip()]
#     h_list = [heading.text for heading in iter_headings(doc.paragraphs)]
#     t1.json = 1
#     t2 = 0
#     t3 = 0
#     t4 = 0
#     # p1 = re.compile(r'\d.*\d', re.S)
#     p1 = p1 = re.compile(r'\d.*\.\d+', re.S)
#     title_list = []
#     for heading in iter_headings(doc.paragraphs):
#         biaoti = ""
#         if heading.style.name == 'Heading 2':
#             t2 = t2 + 1
#             biaoti = str(t1.json) + "." + str(t2)
#             t3 = 0
#         elif heading.style.name == 'Heading 3':
#             t3 = t3 + 1
#             biaoti = str(t1.json) + "." + str(t2) + "." + str(t3)
#             t4 = 0
#         elif heading.style.name == 'Heading 4':
#             t4 = t4 + 1
#             biaoti = str(t1.json) + "." + str(t2) + "." + str(t3) + "." + str(t4)
#         else:
#             continue
#         if not re.findall(p1, heading.text):
#             heading.text = biaoti + heading.text
#         else:
#             heading.text = re.sub(p1, biaoti, heading.text)
#         title_list.append(heading.text)
#     print(title_list)
#     for x in title_list:
#         print(x)
