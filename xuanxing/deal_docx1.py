import docx
from docx import Document
import re

import django

django.setup()
from xuanxing.models import CaseDocx
from xuanxing.models import CaseMiddleType
from xuanxing.models import CaseSamllType

fn = "./t2.docx"
doc = docx.Document(fn)
print(type(doc.paragraphs))
docs = [pj for pj in doc.paragraphs if pj.text.strip()]

tb_dic = {}
te_list = [dc.text for dc in docs if re.match(r"^\d{1}\.\d{1}\.\d{1}.*\D+$", dc.text)]  # 二级标题
te_all_list = [dc.text for dc in docs if re.match(r"^\d{1}\.\d{1}.*\D+$", dc.text)]
te_cha_list = sorted(list(set(te_all_list).difference(set(te_list))))  # 一级标题

for tt in te_list:
    print(tt)
print(len(te_list))
for tec in te_cha_list:
    print(tec)

print(len(doc.tables))
tb_all_list = []
for p_num, paragraph in enumerate(te_list):
    tb_dic = {}
    tb_field_map = {"案例编号": "case_id", "测试目的": "test_goal", "预置条件": "pre_condition", "测试步骤": "test_steps",
                    "预期结果": "expect_result", "实测结果": "test_result", "测试结论": "test_conclusion", "备注": "remark"}
    for tb_num, table in enumerate(doc.tables):
        if p_num == tb_num:
            tb_dic["case_type"] = paragraph
            # tb二级分类
            p_list = [para for para in te_cha_list if para[0:3] == paragraph[0:3]]

            tb_dic["case_middle_type"] = p_list[0]
            tb_dic["case_big_type"] = "数据库"

            for row in table.rows:
                rw_list = [rw.text for rw in row.cells]
                rw_list.pop(0)  # 去除第一个
                rw_list = [rw_list]
                if "操作员" in row.cells[0].text or "复核员" in row.cells[0].text:
                    tb_dic["check_person"] = rw_list
                else:
                    if row.cells[0].text in tb_field_map:
                        if tb_field_map[row.cells[0].text] not in tb_dic:
                            tb_dic[tb_field_map[row.cells[0].text]] = rw_list
                        else:
                            tb_dic[tb_field_map[row.cells[0].text]].extend(rw_list)
            break
    tb_all_list.append(tb_dic)
print(tb_all_list)
for tbs in tb_all_list:
    print(tbs.get("case_type", ""))
    cs_obj = CaseDocx()

    case_lg_type_obj = CaseBigType.objects.filter(l_type_name=tbs.get("case_big_type", ""))
    if case_lg_type_obj:
        cs_obj.case_big_type = case_lg_type_obj[0]
    else:
        # 创建大类
        if tbs.get("case_big_type", ""):
            bg_type_obj = CaseBigType()
            bg_type_obj.l_type_name = tbs.get("case_big_type")
            bg_type_obj.save()
            case_lg_type_obj = CaseBigType.objects.filter(l_type_name=tbs.get("case_big_type", ""))
            cs_obj.case_big_type = case_lg_type_obj[0]
    case_middle_type_obj = CaseMiddleType.objects.filter(m_type_name=tbs.get("case_middle_type", ""))
    if case_middle_type_obj:
        cs_obj.case_middle_type = case_middle_type_obj[0]
    else:
        if tbs.get("case_middle_type", ""):
            md_type_obj = CaseMiddleType()
            md_type_obj.m_type_name = tbs.get("case_middle_type")
            case_lg_type_obj = CaseBigType.objects.filter(l_type_name=tbs.get("case_big_type", ""))
            md_type_obj.l_type = case_lg_type_obj[0]
            md_type_obj.save()
            case_middle_type_obj = CaseMiddleType.objects.filter(m_type_name=tbs.get("case_middle_type", ""))
            cs_obj.case_middle_type = case_middle_type_obj[0]

    case_type_obj = CaseSamllType.objects.filter(s_type_name=tbs.get("case_type", ""))
    print(tbs.get("case_type", ""))
    if case_type_obj:
        cs_obj.case_type = case_type_obj[0]
    else:
        if tbs.get("case_type", ""):

            sm_type_obj = CaseSamllType()
            sm_type_obj.s_type_name = tbs.get("case_type")

            case_md_type_obj = CaseMiddleType.objects.filter(m_type_name=tbs.get("case_middle_type", ""))
            sm_type_obj.m_type = case_md_type_obj[0]
            sm_type_obj.save()
            case_samll_type_obj = CaseSamllType.objects.filter(s_type_name=tbs.get("case_type", ""))
            cs_obj.case_type = case_samll_type_obj[0]

    cs_obj.case_id = tbs.get("case_id", "")
    cs_obj.test_goal = tbs.get("test_goal", "")
    cs_obj.pre_condition = tbs.get("pre_condition", "")
    cs_obj.test_steps = tbs.get("test_steps", "")
    cs_obj.expect_result = tbs.get("expect_result", "")
    cs_obj.test_result = tbs.get("test_result", "")
    cs_obj.test_conclusion = tbs.get("test_conclusion", "")
    cs_obj.remark = tbs.get("remark", "")
    cs_obj.check_person = tbs.get("check_person", "")
    cs_obj.save()


def ttt1():
    tb_dic = {'案例编号': [' ', ' '], '测试目的': ['验证产品是否支持DATABASE/USER', '验证产品是否支持DATABASE/USER'],
              '预置条件': ['数据库集群运行正常', '数据库集群运行正常'],
              '测试步骤': ['创建数据库 prod_test ；\n查看字符集；\n修改字符集；\n创建用户test_user；\n修改用户密码 ；\n删除用户test_user； \n删除数据库prod_test；',
                       '创建数据库 prod_test ；\n查看字符集；\n修改字符集；\n创建用户test_user；\n修改用户密码 ；\n删除用户test_user； \n删除数据库prod_test；'],
              '预期结果': ['成功执行SQL无语法报错', '成功执行SQL无语法报错'],
              '实测结果': ['创建数据库', '符合预期/不支持', '创建用户', '符合预期/不支持', '修改字符集', '符合预期/不支持', '修改密码', '符合预期/不支持', '删除用户',
                       '符合预期/不支持',
                       '删除数据库', '符合预期/不支持'], '测试结论': ['', ''], '备注': ['', ''],
              '操作员：                                复核员：': ['操作员：                                复核员：',
                                                           '操作员：                                复核员：']}

    print(tb_dic)
    document = Document()
    row_num = 0
    col_num = 0
    for k in tb_dic:
        if len(tb_dic[k]) > 0:
            col_num = len(tb_dic[k]) + 1
            break
    row_num = len(tb_dic.keys())
    vvv_list = [len(vvv) for vvv in tb_dic.values()]
    min_vvv = min(vvv_list)
    op_num = sum([(vvn / min_vvv) - 1 for vvn in vvv_list if (vvn / min_vvv) != 1])
    print("*****%s" % op_num)
    row_num += int(op_num)
    print(row_num)
    print(col_num)
    table = document.add_table(rows=int(row_num), cols=int(col_num), style='Table Grid')
    print(tb_dic.items())
    for b, n in enumerate(tb_dic):
        table.cell(b, 0).text = n
        if n != "实测结果":
            for nn, vv in enumerate(tb_dic[n]):
                table.cell(b, nn + 1).text = vv
    # document.save("xxx5.docx")


def create_tb():
    document = docx.Document()
    table = document.add_table(rows=37, cols=13, style='Table Grid')
    table.cell(0, 2).merge(table.cell(2, 2))
    document.save("dssda3.docx")


def create_tss():
    document = Document()
    table = document.add_table(rows=37, cols=13, style='Table Grid')

    document.save('table-1.docx')

    document1 = Document('table-1.docx')
    table = document1.tables[0]
    for row, obj_row in enumerate(table.rows):
        for col, cell in enumerate(obj_row.cells):
            cell.text = cell.text + "%d,%d " % (row, col)

    document1.save('table-2.docx')


def ttt2():
    tb_dic = {'案例编号': [' ', ' '], '测试目的': ['验证产品是否支持DATABASE/USER', '验证产品是否支持DATABASE/USER'],
              '预置条件': ['数据库集群运行正常', '数据库集群运行正常'],
              '测试步骤': ['创建数据库 prod_test ；\n查看字符集；\n修改字符集；\n创建用户test_user；\n修改用户密码 ；\n删除用户test_user； \n删除数据库prod_test；',
                       '创建数据库 prod_test ；\n查看字符集；\n修改字符集；\n创建用户test_user；\n修改用户密码 ；\n删除用户test_user； \n删除数据库prod_test；'],
              '预期结果': ['成功执行SQL无语法报错', '成功执行SQL无语法报错'],
              '实测结果': ['创建数据库', '符合预期/不支持', '创建用户', '符合预期/不支持', '修改字符集', '符合预期/不支持', '修改密码', '符合预期/不支持', '删除用户',
                       '符合预期/不支持',
                       '删除数据库', '符合预期/不支持'], '测试结论': ['', ''], '备注': ['', ''],
              '操作员：                                复核员：': ['操作员：                                复核员：',
                                                           '操作员：                                复核员：']}

    print(tb_dic)
    document = Document()
    row_num = 0
    col_num = 0
    for k in tb_dic:
        if len(tb_dic[k]) > 0:
            col_num = len(tb_dic[k]) + 1
            break
    row_num = len(tb_dic.keys())
    vvv_list = [len(vvv) for vvv in tb_dic.values()]
    row_num = row_num + max(vvv_list) / min(vvv_list) - 1

    print(row_num)
    print(col_num)
    table = document.add_table(rows=int(row_num), cols=int(col_num), style='Table Grid')
    print(tb_dic.items())

# if __name__ == '__main__':
#     # create_tss()
#     ttt1()
