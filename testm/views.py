from django.shortcuts import render

# Create your views here.

from .models import *
from testm.models import *

from django.http import HttpResponse, JsonResponse
from openpyxl import Workbook
from django.shortcuts import render
from docxtpl import DocxTemplate, RichText
from mysite.settings import OUTPUT_ROOT
import docx
from docx.oxml.ns import qn

import tempfile
import io
import zipfile
from wsgiref.util import FileWrapper
from urllib.request import urlretrieve
import datetime
import time
from django.utils.encoding import escape_uri_path
from testm.models import Projects
from django.http import HttpResponse, JsonResponse
from openpyxl import Workbook
from django.shortcuts import render
from tools.tools import juge_user_pj

import os
import openpyxl
import copy


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def Export(report, id, ip):
    """
    下载数据备案单
    :param request:
    :param paper_num: 备案单号
    :return: 数据流
    """
    filename = str(id) + '.docx'
    base_url = OUTPUT_ROOT
    asset_url = base_url + '\\report_template.docx'
    tpl = DocxTemplate(asset_url)
    context = dict()
    context['report'] = report
    tpl.render(context)
    response = HttpResponse(content_type='application/msword')
    # response = HttpResponse(content_type='application/zip')
    response['Content-Disposition'] = 'attachment;filename="{}"'.format(filename)
    tpl.save(response)
    URL = 'http://' + str(ip) + report['test_demand']

    '''#doc.Hyperlinks.Add(Anchor=doc.Range,Address="http://stackoverflow.com/questions/34636391/adding-hyperlinks-in-microsoft-word-using-python")
    URL = 'http://' + str(ip) + report['test_demand']
    res = requests.get(URL)
    res.raise_for_status()
    playFile = open('C:\\Users\\Public\\fujian.docx', 'wb')
    for chunk in res.iter_content(100000):
        print(chunk)
        playFile.write(chunk)

    playFile.close()
    print(playFile.name)

    document = docx.Document("C:\\Users\\guruifeng\\testlink.docx")
    p = document.add_paragraph()
    print(p)
    paragraph = document.add_paragraph()

    document.add_page_break()
    # 添加链接到url
    hyperlink = add_hyperlink(p, "C:\\Users\\guruifeng\\testlink.docx", 'C:\\Users\\guruifeng\\testlink.docx',
                              None, True)

    document.save('demo.docx')'''

    return response


def Export_Report(request, id):
    try:
        if request.method == "GET":
            from django.db import connection
            if request.META.get('HTTP_X_FORWARDED_FOR'):
                ip = request.META.get("HTTP_X_FORWARDED_FOR")
            else:
                ip = request.META.get("REMOTE_ADDR")

            cursor1 = connection.cursor()
            cursor2 = connection.cursor()
            cursor3 = connection.cursor()
            cursor4 = connection.cursor()

            cursor1.execute("select reports_groups_id from report_groups where reports_id=%s", (id,))
            cursor2.execute("select products_id from reports_prods where reports_id=%s", (id,))
            cursor3.execute("select executes_id from report_executes where reports_id=%s", (id,))
            cursor4.execute("select conclusion_id from  products_conclusion where reports_id=%s", (id,))
            rows1 = cursor1.fetchall()
            rows2 = cursor2.fetchall()
            rows3 = cursor3.fetchall()
            rows4 = cursor4.fetchall()

            reports_groups_id = []
            for row in rows1:
                reports_groups_id.append(row[0])
            res0 = Reports_groups.objects.filter(id__in=reports_groups_id)
            reports_groups0 = []
            reports_groups1 = []
            for group in res0:
                groups = dict()
                groups['name'] = group.name
                groups['department'] = group.department
                groups['contacts'] = group.contacts
                groups['responsibilities'] = group.responsibilities
                # groups['group_flag'] = group.group_flag
                if group.group_flag == 0:
                    reports_groups0.append(groups)
                else:
                    reports_groups1.append(groups)
            reports_prods_id = []
            for row3 in rows2:
                reports_prods_id.append(row3[0])
            res0 = Products.objects.filter(id__in=reports_prods_id)
            report_prods = []
            for prod in res0:
                prods = dict()
                prods['id'] = prod.id
                prods['vend_name'] = prod.vend_name
                prods['prod_name'] = prod.prod_name
                prods['technical_index'] = prod.technical_index
                report_prods.append(prods)
            reports_executes_id = []
            for row in rows3:
                reports_executes_id.append(row[0])
            res = Executes.objects.filter(id__in=reports_executes_id)
            execute_time = dict()
            for execute in res:
                time_str = str(execute.start_time).replace('-', '/') + '-' + str(execute.end_time).replace('-', '/')
                if execute.execute_name == '选型启动会':
                    execute_time['选型启动会'] = time_str
                elif execute.execute_name == '选型测试环境准备':
                    execute_time['选型测试环境准备'] = time_str
                elif execute.execute_name == '测试安装调试':
                    execute_time['测试安装调试'] = time_str
                elif execute.execute_name == '测试实施阶段':
                    execute_time['测试实施阶段'] = time_str
                elif execute.execute_name == '测试分析总结':
                    execute_time['测试分析总结'] = time_str
                elif execute.execute_name == '选型结果汇报确认':
                    execute_time['选型结果汇报确认'] = time_str
                else:
                    break
            reports_conclusion_id = []
            for row in rows4:
                reports_conclusion_id.append(row[0])
            res = Conclusion.objects.filter(id__in=reports_conclusion_id)
            report_conclusion = []
            for prod in res:
                prods = dict()
                prods['id'] = prod.id
                res_prod = Products.objects.filter(id=prod.vend_product_id)

                for res_tmp in res_prod:
                    prods['vend_name'] = res_tmp.vend_name
                    prods['prod_name'] = res_tmp.prod_name
                    break
                prods['pass_index'] = prod.pass_index
                prods['evaluate_index'] = prod.evaluate_index
                report_conclusion.append(prods)
            res = Reports.objects.filter(id=id)
            for report in res:
                report_dic = dict()
                report_dic['report_name'] = report.report_name
                report_dic['test_background'] = report.test_background
                report_dic['test_object'] = report.test_object
                if len(str(report.test_demand)) != 0:
                    # report_dic['test_demand'] = playFile.name
                    report_dic['test_demand'] = str(report.test_demand).split('/')[1]
                else:
                    report_dic['test_demand'] = ''
                if len(str(report.pass_index)) != 0:
                    report_dic['pass_index'] = str(report.pass_index).split('/')[1]
                else:
                    report_dic['pass_index'] = ''
                if len(str(report.evaluate_index)) != 0:
                    report_dic['evaluate_index'] = str(report.evaluate_index).split('/')[1]
                else:
                    report_dic['evaluate_index'] = ''
                if len(str(report.test_case)) != 0:
                    report_dic['test_case'] = str(report.test_case).split('/')[1]
                else:
                    report_dic['test_case'] = ''
                from django.utils.safestring import mark_safe

                report_dic['organization_framework'] = RichText(report.organization_framework)

                report_dic['group_part0'] = reports_groups0
                report_dic['group_part1'] = reports_groups1

                report_dic['prod_vend'] = report_prods
                report_dic['test_time'] = execute_time
                report_dic['test_location'] = report.test_location
                report_dic['test_flow'] = RichText(report.test_flow)
                report_dic['test_environment'] = report.test_environment
                report_dic['test_records'] = report.test_records
                report_dic['test_results'] = report.test_results
                report_dic['test_conclusion'] = report_conclusion
                break
            print(report_dic['organization_framework'])
            res1 = Export(report_dic, id, ip)
            return res1

    except Exception as e:
        print(e)
    # return res1


def report_downzip(request, id):
    try:
        if request.method == "GET":
            if request.META.get('HTTP_X_FORWARDED_FOR'):
                ip = request.META.get("HTTP_X_FORWARDED_FOR")
            else:
                ip = request.META.get("REMOTE_ADDR")
            url_str = "http://" + str(ip)
            download_urls = [{"name": "测试报告" + '-' + str(id) + '.docx',
                              "download_url": url_str + "/testm/report/word_down/" + str(id)}]
            res_report = Reports.objects.filter(id=id)
            for res in res_report:
                download_dic = dict()
                print(res.test_demand)
                report_name = res.report_name
                if res.test_demand != '':
                    download_dic = dict()
                    download_dic['name'] = str(res.test_demand).split('/')[1]
                    download_dic['download_url'] = url_str + res.test_demand.url
                    download_urls.append(download_dic)
                    download_urls = copy.deepcopy(download_urls)
                    download_dic.clear()
                if res.pass_index != '':
                    download_dic['name'] = str(res.pass_index).split('/')[1]
                    download_dic['download_url'] = url_str + res.pass_index.url
                    download_urls.append(download_dic)
                    download_urls = copy.deepcopy(download_urls)
                    download_dic.clear()
                if res.evaluate_index != '':
                    download_dic['name'] = str(res.evaluate_index).split('/')[1]
                    download_dic['download_url'] = url_str + res.evaluate_index.url
                    download_urls.append(download_dic)
                    download_urls = copy.deepcopy(download_urls)
                    download_dic.clear()
                if res.test_case != '':
                    download_dic['name'] = str(res.test_case).split('/')[1]
                    download_dic['download_url'] = url_str + res.test_case.url

                    download_urls.append(download_dic)
                    download_urls = copy.deepcopy(download_urls)

            zip_name = report_name + '-' + str(datetime.datetime.now().strftime("%Y-%m-%d")) + '.zip'
            # download_urls = [{"name": "测试报告", "download_url": "http://"+str(ip)+"/testm/report/word_down/22"}, {"name": "test_demand", "download_url": "http://"+str(ip)+"/media/upload/prodcase_23_mNOFWMk.docx"},{"name": "pass_index", "download_url": "http://"+str(ip)+"/media/upload/prodcase_23_mNOFWMk.docx"},{"name": "valuate_index", "download_url": "http://"+ip+"/media/upload/prodcase_23_mNOFWMk.docx"},{"name": "test_case", "download_url": "http://"+str(ip)+"/media/upload/prodcase_23_mNOFWMk.docx"}]
            # download_urls = ['http://127.0.0.1/testm/reports/22', 'http://127.0.0.1/media/upload/prodcase_23_mNOFWMk.docx']
            '''
            download_urls 要批量下载并且压缩的文件
            '''
            # 创建BytesIO
            s = io.BytesIO()
            # 创建一个临时文件夹用来保存下载的文件
            temp = tempfile.TemporaryDirectory()
            # 使用BytesIO生成压缩文件
            zip = zipfile.ZipFile(s, 'w')
            for i in download_urls:
                # f_name = "{}.docx".format(i['name'])
                f_name = i['name']
                local_path = os.path.join(temp.name, f_name)
                # 下载文件
                urlretrieve(i['download_url'], local_path)
                # 把下载文件的写入压缩文件
                # time.sleep(10)
                zip.write(local_path, f_name)
            # 关闭文件
            zip.close()
            # 指针回到初始位置，没有这一句前端得到的zip文件会损坏
            s.seek(0)
            # 用FileWrapper类来迭代器化一下文件对象，实例化出一个经过更适合大文件下载场景的文件对象，实现原理相当与把内容一点点从文件中读取，放到内存，下载下来，直到完成整个下载过程。这样内存就不会担心你一下子占用它那么多空间了。
            wrapper = FileWrapper(s)
            response = HttpResponse(wrapper, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename={}'.format(escape_uri_path(zip_name))

            return response
    except Exception as e:
        return e


# 主页面
def score_tools(request):
    pjname_list = juge_user_pj(request.session.get("user_id"))
    all_pj_list = [Projects.objects.filter(project_name=p_name).first() for p_name in pjname_list]
    pj_name = request.GET.get("pj_name", "")
    if pj_name:
        all_pj_list = [pj_obj for pj_obj in all_pj_list if pj_name in pj_obj.project_name]
    js_res = {"resp": [{"pj_id": pj.id, "pj_name": pj.project_name,
                        "deal_user": " ".join([d_user.username for d_user in pj.deal_user.all()]),
                        "members": " ".join([m_user.username for m_user in pj.members.all()])} for pj in
                       all_pj_list]}

    # rsp = {}
    # if pj_name:
    #     for p in js_res["resp"]:
    #         if pj_name in p.get("pj_name", ""):
    #             rsp = p
    #     if rsp:
    #         js_res = {"resp": [rsp]}
    return render(request, 'score_tools/score_tools.html', js_res)


# 导入并接收文件
def test_zhibiao(request):
    try:
        print('--file upload---')
        resp_file = request.FILES.get('file', "")
        print(resp_file)
        # 项目id
        pj_id = request.POST.get("pj_id", "")
        print(resp_file.name)
        if resp_file.name.split(".")[-1] != "xlsx":
            return JsonResponse({"code": "0005"})
        print(resp_file.size)
        if int(resp_file.size) < 8000:
            return JsonResponse({"code": "0002"})
        print('--------------')
        # 读取前端上传来的文件，并写入服务器目录下
        files_mulu = os.getcwd() + "/static/upload_files/score_zhibiao"
        file_list = os.listdir(files_mulu)
        for fl in file_list:
            if fl.startswith(pj_id):
                print("xxxxxxxxxxxxx")
                rn_file = f"bak{fl}" + ".bak"
                if rn_file in file_list:
                    os.remove(files_mulu + "/" + rn_file)
                os.rename(files_mulu + f"/{fl}", files_mulu + "/" + rn_file)
        with open('static/upload_files/score_zhibiao/' + str(pj_id) + "_" + resp_file.name, 'wb') as f:
            for line in resp_file.chunks():
                f.write(line)
        print("----over---")
        return JsonResponse({"code": "0000"})
    except Exception as e:
        print(e)
        print("-------------file upload  exception----------")
        return JsonResponse({"code": "0001"})


# 展示文件到页面上
def test_show(request):
    if request.method == "POST":
        show_pj_id = request.POST.get("show_pj_id", "")
        if show_pj_id:


            files_mulu = os.getcwd() + "/static/upload_files/score_zhibiao"
            file_list = os.listdir(files_mulu)
            pj_id = show_pj_id[11:]
            show_pj = Projects.objects.filter(id=int(pj_id)).first()
            vend_list = [vend_temp.vend_name + "(%s)" % vend_temp.product_name for vend_temp in
                         show_pj.vend_prod.all()] if show_pj else []
            if len(vend_list) == 0:
                return JsonResponse({"code": "0003"})
            file_flag = 0
            for fl in file_list:
                if fl.startswith(f"daoru{pj_id}"):
                    file_flag = 1
            print(f"file_flag ---{file_flag}")
            if not file_flag:
                return JsonResponse({"code": "0001"})
            else:
                return JsonResponse({"code": "0000"})
        else:
            return JsonResponse({"code": "0002"})
    else:
        try:
            show_pj_id = request.GET.get("show_pj_id", "")
            if show_pj_id:
                files_mulu = os.getcwd() + "/static/upload_files/score_zhibiao"
                file_list = os.listdir(files_mulu)
                pj_id = show_pj_id[11:]  # show_pj_id 是show_result123的形式
                show_pj = Projects.objects.filter(id=int(pj_id)).first()
                vend_list = [vend_temp.vend_name + "(%s)" % vend_temp.product_name for vend_temp in show_pj.vend_prod.all()] if show_pj else []
                show_file = ""
                for fl in file_list:
                    if fl.startswith(f"daoru{pj_id}"):
                        show_file = fl
                print(f"show_file ---{show_file}")
                show_mulu = files_mulu + "/" + show_file  # 目录
                data = openpyxl.load_workbook(show_mulu)  # 加载excel
                pass_table = data["通过性指标"]  # 通过性指标表
                nrows = pass_table.rows  # 通过性指标表的行
                all_list = []
                flg = 0
                for r in nrows:
                    resp_j = {}
                    ll = []
                    for i in r:
                        ll.append(i.value if i.value else "")
                    resp_j["normal_val"] = ll  # 每一行的每个格子里的数据
                    # 第一行显示厂商和产品 剩下的行不显示
                    if not flg:
                        resp_j["calc_val"] = vend_list  # 厂商+产品 名称
                        flg = 1
                    else:
                        resp_j["calc_val"] = ["" for v in vend_list]
                    all_list.append(resp_j)
                resp_json = {"pass_data": all_list}

                # 评价性指标表
                pj_table = data["评价性指标"]
                pj_rows = pj_table.rows  # 评价性表格的每一行
                pj_all_list = []
                pj_flg = 0
                for pj_r in pj_rows:
                    pj_resp = {}
                    pj_ll = []
                    # 每一格的数据
                    for pj_i in pj_r:
                        pj_ll.append(pj_i.value if pj_i.value else "")
                    pj_resp["pj_normal_val"] = pj_ll
                    if not pj_flg:
                        vend_list_x = copy.deepcopy(vend_list)
                        vend_list_x.extend(["pjone"])
                        pj_resp["pj_calc_val"] = vend_list_x
                        pj_flg = 1
                    else:
                        ven_list = ["" for v in vend_list]
                        ven_list.extend(["pjtwo"])
                        pj_resp["pj_calc_val"] = ven_list
                    pj_all_list.append(pj_resp)
                resp_json["pj_data"] = pj_all_list
                resp_json["test_pj_id"] = pj_id
                return render(request, 'score_tools/score_show.html', resp_json)
        except Exception as e:
            print(e)
            return HttpResponse("文件错误，请重新导入指标文件")


# 清除文件
def test_clear(request):
    try:
        claer_pj_id = request.POST.get("claer_pj_id", "")
        if claer_pj_id:
            files_mulu = os.getcwd() + "/static/upload_files/score_zhibiao"
            file_list = os.listdir(files_mulu)
            pj_id = claer_pj_id[12:]
            for fl in file_list:
                if fl.startswith(f"daoru{pj_id}"):
                    rn_file = f"bak{fl}" + ".bak"
                    if rn_file in file_list:
                        os.remove(files_mulu + "/" + rn_file)  # 清除导入的指标文件的备份文件
                    os.rename(files_mulu + f"/{fl}", files_mulu + "/" + rn_file)  # 对现在的导入的指标文件更名
            result_flag = 0
            for fs in file_list:
                if fs.startswith(f"test_result_{pj_id}"):
                    fs_file = f"bak{fs}" + ".bak"
                    if fs_file in file_list:
                        os.remove(files_mulu + "/" + fs_file)
                    os.rename(files_mulu + f"/{fs}", files_mulu + "/" + fs_file)
                    result_flag = 1
            if not result_flag:
                return JsonResponse({"code": "0001"})

            return JsonResponse({"code": "0000"})
        else:
            return JsonResponse({"code": "0002"})
    except Exception as e:
        print(e)
        return JsonResponse({"code": "0003"})



# 保存文件
def test_save(request):
    save_pj_id = request.POST.get("save_pj_id")
    save_type = request.POST.get("save_type")
    save_ps_data = request.POST.get("save_ps_data")
    if save_ps_data:
        save_ps_data = eval(save_ps_data)
    file_name = f"test_result_{save_pj_id}" + ".xlsx"
    files_mulu = os.getcwd() + "/static/upload_files/score_zhibiao"
    if "tg" == save_type:
        file_list = os.listdir(files_mulu)
        if file_name not in file_list:
            wb = Workbook()
            print(save_ps_data)
            print(type(save_ps_data))
            wb.create_sheet("通过性指标")
            ws = wb.get_sheet_by_name("通过性指标")
            for sv in save_ps_data:
                print(sv)
                ws.append(sv)
            wb.save(files_mulu + "/" + file_name)
            return JsonResponse({"code": "0000"})
        else:
            show_mulu = files_mulu + "/" + file_name
            sheet_obj = openpyxl.load_workbook(show_mulu)
            tg_sheet = "通过性指标"
            if tg_sheet in sheet_obj.sheetnames:
                sheet_obj.remove(sheet_obj[tg_sheet])
            sheet_obj.create_sheet(tg_sheet)
            ws = sheet_obj.get_sheet_by_name(tg_sheet)
            for sv in save_ps_data:
                print(sv)
                ws.append(sv)
            sheet_obj.save(show_mulu)
            return JsonResponse({"code": "0000"})

    if "pj" == save_type:
        rank_data = request.POST.get("save_score_rnak", {})
        print("--------11111-------")
        if rank_data:
            rank_data = eval(rank_data)
        print(rank_data)
        sort_rank_data = sorted(rank_data.items(), key=lambda item: item[1], reverse=True)
        print(sort_rank_data)

        file_list = os.listdir(files_mulu)
        if file_name not in file_list:
            wb = Workbook()
            print(save_ps_data)
            print(type(save_ps_data))
            wb.create_sheet("评价性指标")
            ws = wb.get_sheet_by_name("评价性指标")
            for sv in save_ps_data:
                print(sv)
                ws.append(sv)
            wb.create_sheet("厂商评分")
            wq = wb.get_sheet_by_name("厂商评分")
            for wq_item in sort_rank_data:
                wq.append(wq_item)
            wb.save(files_mulu + "/" + file_name)
            return JsonResponse({"code": "0000"})
        else:
            show_mulu = files_mulu + "/" + file_name
            sheet_obj = openpyxl.load_workbook(show_mulu)
            pj_sheet = "评价性指标"
            if pj_sheet in sheet_obj.sheetnames:
                sheet_obj.remove(sheet_obj[pj_sheet])
            sheet_obj.create_sheet(pj_sheet)
            vend_sheet = "厂商评分"
            if vend_sheet in sheet_obj.sheetnames:
                sheet_obj.remove(sheet_obj[vend_sheet])
            sheet_obj.create_sheet(vend_sheet)
            # 存入评价表
            ws = sheet_obj.get_sheet_by_name(pj_sheet)
            for sv in save_ps_data:
                print(sv)
                ws.append(list(map(lambda x: x.strip(), sv)))
            # 存入厂商评分表
            vends = sheet_obj.get_sheet_by_name(vend_sheet)
            for vd in sort_rank_data:
                vends.append(vd)
            sheet_obj.save(show_mulu)
            return JsonResponse({"code": "0000"})


# 导出文件
def test_export(request):
    export_pj_id = request.GET.get("export_pj_id", "")
    if export_pj_id:
        print(export_pj_id)
        files_mulu = os.getcwd() + "/static/upload_files/score_zhibiao"
        file_list = os.listdir(files_mulu)
        pj_id = export_pj_id[13:]
        for fl in file_list:
            if fl.startswith(f"test_result_{pj_id}"):
                wb = openpyxl.load_workbook(files_mulu + "/" + f"test_result_{pj_id}.xlsx")
                if "Sheet" in wb.sheetnames:
                    wb.remove(wb["Sheet"])
                response = HttpResponse(content_type='application/msexcel')
                response['Content-Disposition'] = "attachment; filename*=utf-8''{}".format(
                    escape_uri_path(f"项目{pj_id}测试结果.xlsx"))
                # response['Content-Disposition'] = 'attachment; filename=' + f"test_result_{pj_id}" + '.xlsx'
                wb.save(response)
                return response
        return HttpResponse("文件不存在")


# 导出模板
def muban_download(request):
    files_mulu = os.getcwd() + "/static/upload_files"
    show_mulu = files_mulu + "/" + "打分模板.xlsx"
    wb = openpyxl.load_workbook(show_mulu)
    response = HttpResponse(content_type='application/msexcel')
    response['Content-Disposition'] = "attachment; filename*=utf-8''{}".format(escape_uri_path("模板.xlsx"))
    # response['Content-Disposition'] = 'attachment; filename=' + 'muban' + '.xlsx'
    wb.save(response)
    return response


# 厂商排序
def vend_rank(request):
    files_mulu = os.getcwd() + "/static/upload_files/score_zhibiao"
    file_list = os.listdir(files_mulu)
    rank_pj_id = request.POST.get("rank_pj_id", "")
    pj_id = rank_pj_id[9:]
    file_flag = 0
    rtn_str = ""
    for fl in file_list:
        if fl.startswith(f"test_result_{pj_id}"):
            show_mulu = files_mulu + "/" + f"test_result_{pj_id}" + ".xlsx"
            data = openpyxl.load_workbook(show_mulu)  # 加载excel
            names = data.sheetnames
            # 通过性指标表
            if "厂商评分" not in data:
                return JsonResponse({"code": "0010"})
            rank_data = data["厂商评分"]
            nrows = rank_data.rows
            rank_list = []
            for rank_row in nrows:
                for r in rank_row:
                    rtn_str += r.value if r.value else ""
                    rtn_str += "&nbsp;&nbsp;&nbsp;&nbsp;"
                rtn_str += "<br>"
            file_flag = 1

    if not file_flag:
        return JsonResponse({"code": "0001"})
    else:
        return JsonResponse({"code": "0000", "rank_data": rtn_str})
